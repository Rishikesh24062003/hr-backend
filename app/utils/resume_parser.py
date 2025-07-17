import os
import re
import logging
import pdfplumber
from datetime import datetime
from PyPDF2 import PdfReader
from docx import Document

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

def extract_text_from_pdf(file_path):
    """Enhanced PDF text extraction with fallback options."""
    try:
        text = ""
        
        # First attempt with PyPDF2
        try:
            with open(file_path, 'rb') as file:
                reader = PdfReader(file)
                
                # Check if PDF is encrypted
                if reader.is_encrypted:
                    logger.warning(f"PDF is encrypted: {file_path}")
                    return None
                
                # Extract text from all pages
                for page_num, page in enumerate(reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
                    except Exception as e:
                        logger.warning(f"Error extracting text from page {page_num + 1}: {str(e)}")
                        continue
                        
        except Exception as e:
            logger.error(f"PyPDF2 extraction failed for {file_path}: {str(e)}")
            # Try fallback method with pdfplumber if available
            try:
                
                text = ""
                with pdfplumber.open(file_path) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
            except ImportError:
                logger.error("pdfplumber not available for fallback")
                return None
            except Exception as fallback_error:
                logger.error(f"Fallback extraction also failed: {str(fallback_error)}")
                return None
        
        # Check if any text was extracted
        if not text or not text.strip():
            logger.error(f"No text extracted from PDF: {file_path}")
            return None
            
        return text.strip()
        
    except Exception as e:
        logger.error(f"Failed to extract text from PDF {file_path}: {str(e)}")
        return None

def extract_text_from_docx(file_path):
    """Enhanced DOCX text extraction with error handling."""
    try:
        doc = Document(file_path)
        text_parts = []
        
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text.strip())
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text.strip())
        
        text = "\n".join(text_parts)
        
        # Check if any text was extracted
        if not text or not text.strip():
            logger.error(f"No text extracted from DOCX: {file_path}")
            return None
            
        return text.strip()
        
    except Exception as e:
        logger.error(f"Failed to extract text from DOCX {file_path}: {str(e)}")
        return None

def extract_text_from_file(file_path):
    """Enhanced file text extraction with comprehensive validation."""
    try:
        # Validate file exists and has content
        if not os.path.exists(file_path):
            logger.error(f"File does not exist: {file_path}")
            return None
        
        file_size = os.path.getsize(file_path)
        if file_size == 0:
            logger.error(f"File is empty: {file_path}")
            return None
        
        # Get file extension
        file_extension = os.path.splitext(file_path)[1].lower()
        
        # Extract text based on file type
        if file_extension == '.pdf':
            return extract_text_from_pdf(file_path)
        elif file_extension in ['.docx', '.doc']:
            return extract_text_from_docx(file_path)
        else:
            logger.error(f"Unsupported file format: {file_extension}")
            return None
            
    except Exception as e:
        logger.error(f"Unexpected error in text extraction: {str(e)}")
        return None

def parse_resume(file_path):
    """Enhanced resume parsing with comprehensive error handling."""
    try:
        # Extract raw text
        raw_text = extract_text_from_file(file_path)
        logger.info(f"Extracted {len(raw_text or '')} characters")
        
        if not raw_text:
            raise Exception("Could not extract text from resume file")
        
        # Continue with existing parsing logic...
        structured_data = {
            'name': extract_name(raw_text),
            'email': extract_email(raw_text),
            'phone': extract_phone(raw_text),
            'skills': extract_skills(raw_text),
            'experience_years': calculate_experience_years(raw_text),
            'education': extract_education(raw_text),
        }
        
        return {
            'raw_text': raw_text,
            'structured_data': structured_data,
            'name': structured_data['name'],
            'email': structured_data['email'],
            'phone': structured_data['phone'],
            'parsing_timestamp': datetime.utcnow().isoformat()
        }
        
    except Exception as e:
        logger.error(f"Failed to parse resume {file_path}: {str(e)}")
        raise Exception(f"Resume parsing failed: {str(e)}")

# Include your existing helper functions (extract_name, extract_email, etc.)

def extract_name(text):
    """Placeholder for name extraction."""
    # Basic regex for a common name pattern (e.g., two words at the beginning)
    match = re.search(r"^[A-Z][a-z]+(?:\s[A-Z][a-z]+)*", text)
    return match.group(0) if match else "Unknown Name"

def extract_email(text):
    """Placeholder for email extraction."""
    match = re.search(r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}", text)
    return match.group(0) if match else "Unknown Email"

def extract_phone(text):
    """Placeholder for phone number extraction."""
    match = re.search(r"\b\d{3}[-.]?\d{3}[-.]?\d{4}\b", text)
    return match.group(0) if match else "Unknown Phone"

def extract_skills(text):
    """Placeholder for skill extraction."""
    # Example: Look for common tech skills (very basic)
    skills_keywords = ["Python", "JavaScript", "React", "SQL", "Flask", "Machine Learning"]
    found_skills = [skill for skill in skills_keywords if re.search(r"\b" + re.escape(skill) + r"\b", text, re.IGNORECASE)]
    return found_skills

def calculate_experience_years(text):
    """Placeholder for experience calculation."""
    # This is a very simplistic placeholder
    match = re.search(r"(\d+)\s*years?\s*experience", text, re.IGNORECASE)
    if match:
        try:
            return int(match.group(1))
        except ValueError:
            pass
    return 0 # Default to 0 years if not found or parse error

def extract_education(text):
    """Placeholder for education extraction."""
    # Example: Look for degree patterns
    education_keywords = ["Bachelor", "Master", "Ph.D.", "University", "College"]
    found_education = [edu for edu in education_keywords if re.search(r"\b" + re.escape(edu) + r"\b", text, re.IGNORECASE)]
    return found_education
